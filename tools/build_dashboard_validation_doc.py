from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Volumes/PortableSSD/newtaskdashboard")
OUT = ROOT / "docs" / "업무현황_대시보드_검증서_2026-07-15.docx"
FORGE_SHOT = Path("/tmp/dashboard-forge-completion.png")
PEOPLE_SHOT = Path("/tmp/dashboard-people-tasks.png")
MAGIC_SHOT = Path("/tmp/magic-cauldron-dashboard-final.png")

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
TEXT = "202A35"
MUTED = "64748B"
LIGHT_BLUE = "EAF2F8"
LIGHT_GRAY = "F2F4F7"
GREEN = "198754"
LIGHT_GREEN = "EAF7EF"
AMBER = "A15C00"
LIGHT_AMBER = "FFF4DB"
RED = "B42318"
LIGHT_RED = "FDECEC"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_inches, indent_twips=0):
    total = int(sum(widths_inches) * 1440)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_twips))
    tbl_ind.set(qn("w:type"), "dxa")

    old_grid = table._tbl.tblGrid
    for child in list(old_grid):
        old_grid.remove(child)
    for width in widths_inches:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(width * 1440)))
        old_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx >= len(widths_inches):
                continue
            width_twips = int(widths_inches[idx] * 1440)
            cell.width = Inches(widths_inches[idx])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width_twips))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_repeat_table_rows_disabled(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_run_font(run, size=None, bold=None, color=None, name="AppleGothic"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_hyperlink(paragraph, text, url, color=BLUE, underline=True):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "AppleGothic")
    r_fonts.set(qn("w:hAnsi"), "AppleGothic")
    r_fonts.set(qn("w:eastAsia"), "AppleGothic")
    r_pr.append(r_fonts)
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    r_pr.append(c)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        r_pr.append(u)
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, 9, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def keep_paragraph(paragraph, keep_next=False):
    p_pr = paragraph._p.get_or_add_pPr()
    keep_lines = OxmlElement("w:keepLines")
    p_pr.append(keep_lines)
    if keep_next:
        node = OxmlElement("w:keepNext")
        p_pr.append(node)
    widow = OxmlElement("w:widowControl")
    p_pr.append(widow)


def add_title(doc, text, subtitle=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(26)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_run_font(run, 26, True, DARK_BLUE)
    keep_paragraph(p, True)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(18)
        r = p2.add_run(subtitle)
        set_run_font(r, 12, False, MUTED)
        keep_paragraph(p2, True)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    keep_paragraph(p, True)
    return p


def add_body(doc, text="", bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, 11, True, TEXT)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, 11, False, TEXT)
    else:
        r = p.add_run(text)
        set_run_font(r, 11, False, TEXT)
    keep_paragraph(p)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.left_indent = Inches(0.5 + 0.25 * level)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    set_run_font(r, 11, False, TEXT)
    keep_paragraph(p)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    set_run_font(r, 11, False, TEXT)
    keep_paragraph(p)
    return p


def add_callout(doc, title, body, fill=LIGHT_BLUE, title_color=DARK_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run_font(r, 11, True, title_color)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_run_font(r2, 10.5, False, TEXT)
    return table


def add_table(doc, headers, rows, widths, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, text in enumerate(headers):
        cell = hdr.cells[idx]
        set_cell_shading(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_run_font(r, font_size, True, DARK_BLUE)
    for row_data in rows:
        row = table.add_row()
        set_repeat_table_rows_disabled(row)
        for idx, value in enumerate(row_data):
            cell = row.cells[idx]
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            color = TEXT
            if str(value) == "통과":
                color = GREEN
            elif str(value) in {"조건부 통과", "부분 통과", "운영 검증 필요"}:
                color = AMBER
            elif str(value) in {"실패", "미통과"}:
                color = RED
            set_run_font(r, font_size, str(value) in {"통과", "조건부 통과", "부분 통과", "운영 검증 필요", "실패", "미통과"}, color)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_figure(doc, path, caption, width=6.25):
    if not path.exists():
        add_callout(doc, "이미지 누락", str(path), LIGHT_RED, RED)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    r = cap.add_run(caption)
    set_run_font(r, 9, False, MUTED)
    keep_paragraph(cap)


def add_page_break(doc):
    doc.add_page_break()


def build_document():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "AppleGothic"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "AppleGothic")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "AppleGothic")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "AppleGothic")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Title", 26, DARK_BLUE, 0, 8),
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = "AppleGothic"
        style._element.rPr.rFonts.set(qn("w:ascii"), "AppleGothic")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "AppleGothic")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "AppleGothic")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_style in ("List Bullet", "List Bullet 2", "List Number"):
        style = styles[list_style]
        style.font.name = "AppleGothic"
        style._element.rPr.rFonts.set(qn("w:ascii"), "AppleGothic")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "AppleGothic")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "AppleGothic")
        style.font.size = Pt(11)

    for sec in doc.sections:
        hp = sec.header.paragraphs[0]
        hp.text = "몰입 업무현황 대시보드  |  검증 문서"
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in hp.runs:
            set_run_font(run, 8.5, False, MUTED)
        add_page_number(sec.footer.paragraphs[0])

    props = doc.core_properties
    props.title = "업무현황 의사결정 대시보드 검증서"
    props.subject = "의사결정 지연과 팀 대기시간 감소를 위한 기능·데이터·운영 검증"
    props.author = "Codex"
    props.keywords = "업무현황, 의사결정, 프로젝트, 병목, Notion, Slack, GPT"

    add_title(
        doc,
        "업무현황 의사결정\n대시보드 검증서",
        "의사결정 지연과 팀 전체 대기시간 감소를 위한 기능·데이터·운영 검증",
    )
    meta = add_table(
        doc,
        ["구분", "내용"],
        [
            ["검증 대상", "decision-dashboard prototype v0.9 / finalized.html"],
            ["기준 일자", "2026-07-15 (KST)"],
            ["검증 범위", "대표 브리핑 · 프로젝트 · 담당별 흐름 · 확인 필요"],
            ["문서 상태", "조건부 검증 가능 — 운영 파이프라인은 별도 검증 필요"],
        ],
        [1.35, 5.15],
        10,
    )
    add_callout(
        doc,
        "검증 결론",
        "프로토타입은 ‘프로젝트가 어떻게 진행되는가, 얼마나 진행됐는가, 지금 어디가 막혔는가’를 확인하는 화면 구조를 갖췄다. 다만 매일 오전 자동 수집, 실제 Slack·회의록 추출, GPT 요약의 사실성, 의존관계 커버리지는 정적 HTML만으로 검증할 수 없으므로 구현 완료 조건에 포함해야 한다.",
        LIGHT_AMBER,
        AMBER,
    )
    add_heading(doc, "문서 사용법", 2)
    add_number(doc, "검토자는 먼저 1장의 출시 판단과 2장의 판정 원칙을 확인한다.")
    add_number(doc, "3~5장에서 실제 화면과 사례 데이터가 한 번에 이해되는지 확인한다.")
    add_number(doc, "6장의 검증 시나리오를 실행하고 8장의 서명란에 결과를 남긴다.")
    add_body(doc, "이 문서는 설계 의도와 현재 프로토타입의 실제 검증 결과를 구분한다. ‘통과’는 지금 확인된 사항, ‘운영 검증 필요’는 구현·실데이터 연동 후 확인해야 하는 사항이다.")

    add_page_break(doc)
    add_heading(doc, "1. 검증 목적과 최종 판단", 1)
    add_heading(doc, "1.1 시스템 목표", 2)
    add_body(doc, "이 시스템의 목적은 업무를 많이 관리하거나 개인을 평가하는 것이 아니다. 프로젝트 진행 상태와 계획 대비 진척, 의사결정 지연, 파트 간 병목을 조기에 발견해 팀 전체의 대기시간을 줄이는 것이다.")
    add_bullet(doc, "대표가 ‘오늘 결정할 것 → 현재 막힌 것 → 어제와 달라진 것’ 순서로 실행 우선순위를 이해한다.")
    add_bullet(doc, "프로젝트별로 핵심 스펙의 확정 여부, 세부 일감, 완료율, 파트 간 의존관계와 후속 영향을 함께 본다.")
    add_bullet(doc, "개인 탭은 작업량 평가가 아니라 누가 누구의 결정·산출물을 기다리는지, 어떤 업무가 팀 대기를 만드는지 보여준다.")
    add_bullet(doc, "Notion·Slack·회의록이 충돌하면 출처별 주장을 함께 보이고 ‘확인 필요’로 올린다.")
    add_bullet(doc, "대시보드는 읽기 전용이다. 대표의 결정을 저장하거나 업무 원본을 자동으로 덮어쓰지 않는다.")

    add_heading(doc, "1.2 핵심 검증 질문", 2)
    questions = [
        "프로젝트가 지금 어떤 단계이며 목표일까지 얼마나 진행됐는가?",
        "어떤 핵심 스펙이 미확정·지연되어 다음 파트의 대기를 만들고 있는가?",
        "같은 사실에 대해 Notion·Slack·회의록이 다르게 말하고 있지는 않은가?",
        "누가 많은 일을 보유했는가보다, 누구의 어떤 일감 때문에 팀 흐름이 멈추는가?",
        "요약만 읽어도 원문을 추가로 뒤지지 않고 다음 확인·결정을 정할 수 있는가?",
    ]
    for q in questions:
        add_bullet(doc, q)

    add_heading(doc, "1.3 출시·구현 판단", 2)
    add_callout(
        doc,
        "조건부 가능",
        "화면 구조와 핵심 데이터 표현은 검증 가능 수준이다. 그러나 운영 출시 전에는 ① 오전 자동 수집, ② Slack·회의록 추출의 원문 추적성, ③ GPT 요약 스키마와 오류 처리, ④ 의존관계 등록률, ⑤ 충돌 확인 처리량을 반드시 실데이터로 검증해야 한다.",
        LIGHT_GREEN,
        GREEN,
    )

    add_page_break(doc)
    add_heading(doc, "2. 데이터와 판정 원칙", 1)
    add_heading(doc, "2.1 수집·요약 워크플로우", 2)
    flow_rows = [
        ["1. 대상 선택", "프로젝트 리스트에서 ‘요약’이 체크된 프로젝트만 수집 대상에 포함"],
        ["2. 매일 오전 수집", "목표 운영안: 매일 08:00 KST에 스케줄러/통합 워커가 Notion·Slack·회의록을 읽음"],
        ["3. 구조화", "스펙–자식 일감 관계, 담당자, 상태, 기간, 출처, 원문 링크를 정규화"],
        ["4. 규칙 판정", "완료율·기한·스펙 확정·의존성·후속 영향·출처 충돌을 규칙으로 계산"],
        ["5. GPT 요약", "규칙 결과와 출처가 있는 사실만 문장화. 새 상태·담당·기한을 추측하지 않음"],
        ["6. 화면 제공", "대표 브리핑 → 프로젝트 → 담당별 흐름 → 확인 필요 순으로 읽기 전용 제공"],
    ]
    add_table(doc, ["단계", "설계 내용"], flow_rows, [1.35, 5.15], 9.5)
    add_callout(
        doc,
        "중요: 자동 수집의 담당 주체",
        "사람이 매일 수동으로 갱신하는 것이 아니라 스케줄러가 통합 수집 작업을 실행하고, 실패 시 운영 담당자에게 알림을 보내는 구조가 목표다. 현재 정적 프로토타입에는 스케줄러와 실패 알림이 구현되어 있지 않으므로 운영 검증 항목이다.",
        LIGHT_AMBER,
        AMBER,
    )

    add_heading(doc, "2.2 출처별 역할", 2)
    source_rows = [
        ["Notion", "상태 판정의 기준 원장", "스펙, 자식 일감, 담당, 상태, 기간, 계층"],
        ["Slack", "운영 신호·충돌 탐지", "실제 협의, 일정 변경 언급, 대기·막힘·질문"],
        ["회의록", "결정·합의 맥락", "결정 사항, 보류, 책임자, 후속 조치"],
        ["GPT", "사실 기반 문장화", "세 출처의 근거를 묶어 짧은 브리핑 생성"],
    ]
    add_table(doc, ["출처", "주요 역할", "사용 범위"], source_rows, [1.05, 1.8, 3.65], 9.2)
    add_body(doc, "Slack 운영 신호는 별도 독립 메뉴가 아니라 프로젝트 요약과 ‘확인 필요’에 흡수한다. 충돌이 없는 신호는 프로젝트 상황 설명에, 출처 간 상태·기한·담당 충돌은 확인 필요에 표시한다.")

    add_heading(doc, "2.3 완료율과 프로젝트 건강상태", 2)
    add_body(doc, "완료율 = 완료된 자식 일감 수 ÷ 해당 스펙의 전체 자식 관계 수 × 100")
    add_bullet(doc, "완료율은 진행 참고값이며 프로젝트 건강상태와 동일하지 않다.")
    add_bullet(doc, "프로젝트 건강상태는 핵심 스펙 확정 여부, 파트 간 의존관계, 기한과 후속 작업 영향으로 판단한다.")
    add_bullet(doc, "스펙 상세·완료율은 부모 스펙에 연결된 전체 자식 관계를 사용한다.")
    add_bullet(doc, "담당별 현재 업무량은 완료 일감을 제외한 현재 프로젝트 뷰를 사용한다.")
    add_bullet(doc, "삭제·종료된 자식이 진행 중으로 남아 있으면 완료율에 조용히 반영하지 않고 데이터 오류로 표시한다.")

    add_heading(doc, "2.4 GPT 요약 명령 계약", 2)
    add_body(doc, "운영 구현에서는 아래 지침을 시스템 프롬프트와 출력 스키마에 고정한다.")
    for item in [
        "입력에 있는 사실만 사용하고, 상태·담당·기한·완료율을 추측하지 않는다.",
        "프로젝트마다 ‘현재 단계 / 완료율 / 막힌 원인 / 후속 영향 / 확인 필요 / 근거 출처’를 출력한다.",
        "출처 충돌 시 어느 쪽이 최신인지 임의 판단하지 않고 각 주장과 시각·링크를 함께 제시한다.",
        "근거가 부족하면 ‘정보 부족’으로 표시하고 긍정적·부정적 상태를 단정하지 않는다.",
        "요약 우선순위는 ‘오늘 결정할 것 → 막힌 것 → 어제와 달라진 것’으로 유지한다.",
        "개인 요약은 평가 문구를 금지하고, 기다리는 상대·대기 원인·영향받는 작업만 기술한다.",
        "JSON Schema 검증에 실패하면 GPT 문장을 폐기하고 규칙 엔진의 사실 목록 템플릿으로 대체한다.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "3. 화면별 검증 기준", 1)
    screen_rows = [
        ["대표 브리핑", "오늘 결정할 것 → 막힌 것 → 어제와 달라진 것", "각 항목에 프로젝트·스펙·영향·근거가 있어야 함"],
        ["프로젝트", "요약 체크된 모든 프로젝트와 핵심 스펙", "스펙 토글에서 세부 일감·완료율·출처·충돌 확인"],
        ["담당별 흐름", "현재 뷰에 있는 전체 인원", "토글에서 프로젝트·상위 스펙·일감·상태·기간·링크 확인"],
        ["확인 필요", "핵심 필드 충돌과 데이터 오류", "출처별 내용을 덮어쓰지 않고 나란히 표시"],
    ]
    add_table(doc, ["화면", "반드시 보여줄 내용", "검증 포인트"], screen_rows, [1.1, 2.45, 2.95], 9.0)

    add_heading(doc, "3.1 현재 프로토타입 데이터 범위", 2)
    scope_rows = [
        ["요약 대상 프로젝트", "2개", "피자레디, 포지 앤 포춘"],
        ["최상위 스펙", "52개", "피자레디 14, 포지 앤 포춘 38"],
        ["현재 뷰 자식 일감", "112개", "피자레디 24, 포지 앤 포춘 88"],
        ["확인된 담당자", "15명", "77개 담당 연결"],
        ["미배정 팀 큐", "65개", "개인 평가가 아니라 팀 대기·배정 위험으로 표시"],
    ]
    add_table(doc, ["항목", "수량", "해석"], scope_rows, [1.55, 0.8, 4.15], 9.4)

    add_heading(doc, "3.2 대표가 한 번에 이해해야 하는 답", 2)
    add_callout(
        doc,
        "예시 결론",
        "포지 앤 포춘은 현재 뷰 기준 88개 일감 중 미배정 57개와 기한 초과 12개가 우선 위험이다. 피자레디 ‘마법가마솥’은 자식 14개 중 7개 완료(50%)지만, Notion 7/31과 Slack 약 8/7의 일정 주장이 충돌하고 SP61 제작 시작이 대기 중이므로 완료율만으로 정상 판정하면 안 된다.",
        LIGHT_BLUE,
        DARK_BLUE,
    )

    add_heading(doc, "4. 실제 데이터 사례 검증", 1)
    add_heading(doc, "4.1 피자레디 — 라이브이벤트 ‘마법가마솥’", 2)
    magic_rows = [
        ["전체 자식", "14개", "부모 스펙 전체 관계 기준"],
        ["완료", "7개", "완료율 50%"],
        ["시작 전", "4개", "후속 제작 대기 포함"],
        ["계획", "2개", "담당·기간 미지정"],
        ["데이터 오류", "1개", "삭제된 자식이 진행 중으로 남음"],
        ["현재 단계", "기획 정리", "SP61 제작 시작 대기"],
        ["일정 충돌", "Notion 7/31 / Slack 약 8/7", "자동 선택하지 않고 확인 필요"],
    ]
    add_table(doc, ["항목", "현재 값", "의사결정 해석"], magic_rows, [1.35, 1.85, 3.3], 9.4)
    add_body(doc, "검토자가 내려야 할 판단: 완료율은 50%이지만 일정 기준이 충돌하고 후속 제작이 시작되지 않았으므로 ‘정상 진행’으로 단정하지 않는다. 일정 기준을 확인하고 SP61 제작 착수 조건을 명확히 해야 한다.")
    p = doc.add_paragraph()
    r = p.add_run("근거: ")
    set_run_font(r, 10, True, TEXT)
    add_hyperlink(p, "Notion 스펙", "https://app.notion.com/p/35eb4a46500380ddb0dde7d8fc52d438")
    r2 = p.add_run("  |  ")
    set_run_font(r2, 10, False, MUTED)
    add_hyperlink(p, "Slack 일정 스레드", "https://molip0225.slack.com/archives/C05FT5754SK/p1781258173365049?thread_ts=1781258173.365049&cid=C05FT5754SK")
    add_figure(doc, MAGIC_SHOT, "그림 1. 피자레디 ‘마법가마솥’ 스펙 상세와 출처 충돌 표시", 5.3)

    add_page_break(doc)
    add_heading(doc, "4.2 포지 앤 포춘 — 우선 스펙 완료율", 2)
    forge_rows = [
        ["Spinner", "0 / 6", "0%", "착수 전"],
        ["Sprint 2 개선", "1 / 7", "14%", "초기 진행"],
        ["제련소·분류기", "4 / 10", "40%", "기본 개발 완료, 응답·후속 대기"],
        ["시식코너", "0 / 6", "0%", "착수 전"],
        ["07/13 일감", "1 / 31", "3%", "확인 요청 17개가 주요 큐"],
        ["UI 전환", "1 / 7", "14%", "초기 진행"],
    ]
    add_table(doc, ["스펙", "완료", "완료율", "현재 해석"], forge_rows, [1.55, 0.8, 0.75, 3.4], 9.2)
    add_body(doc, "검토자가 내려야 할 판단: ‘07/13 일감’은 완료율 3%라는 숫자보다 17개 확인 요청이 제작 흐름을 막는다는 점이 우선이다. ‘제련소·분류기’는 화면의 현재 뷰 6개가 아니라 전체 자식 관계 10개를 기준으로 40%를 계산한다.")
    add_figure(doc, FORGE_SHOT, "그림 2. 포지 앤 포춘 스펙별 완료율과 토글 상세", 6.1)

    add_page_break(doc)
    add_heading(doc, "4.3 담당별 흐름 — 전체 인원과 실제 일감", 2)
    people_rows = [
        ["Kady", "15"], ["Ink", "13"], ["Hati", "9"], ["Momo", "7"], ["Likey", "6"],
        ["Modin", "5"], ["Lua", "4"], ["Hongki", "3"], ["Trisa", "3"], ["Jiten", "3"],
        ["Sian", "3"], ["Hank", "2"], ["Wade", "2"], ["담당 미상", "1"], ["Search", "1"],
    ]
    compact = []
    for i in range(0, len(people_rows), 3):
        chunk = people_rows[i:i+3]
        row = []
        for pair in chunk:
            row.extend(pair)
        while len(row) < 6:
            row.extend(["", ""])
        compact.append(row)
    add_table(doc, ["담당", "일감", "담당", "일감", "담당", "일감"], compact, [1.25, 0.6, 1.25, 0.6, 1.25, 0.6], 9.4)
    add_body(doc, "15명 모두 토글로 열 수 있으며, 각 항목은 프로젝트·상위 스펙·자식 일감·상태·기간·Notion 링크를 표시한다. 수량은 완료 일감을 제외한 현재 프로젝트 뷰 기준이다.")
    add_callout(
        doc,
        "해석 주의",
        "15개를 보유한 Kady를 자동으로 ‘문제 인력’으로 평가하지 않는다. 다수 일감이 같은 확인·승인·산출물에 묶여 팀 대기를 만들고 있는지, 미배정 65개가 특정 역할로 몰리는지를 확인하는 출발점으로 사용한다.",
        LIGHT_AMBER,
        AMBER,
    )
    add_figure(doc, PEOPLE_SHOT, "그림 3. 담당자 토글에서 실제 현재 일감을 펼친 화면", 6.1)

    add_page_break(doc)
    add_heading(doc, "5. 검증 시나리오와 현재 결과", 1)
    test_rows = [
        ["V-01", "요약 체크 프로젝트 선택", "체크된 모든 프로젝트만 표시", "피자레디·포지 앤 포춘 2개 확인", "통과"],
        ["V-02", "매일 08:00 자동 수집", "스케줄러 실행·실패 알림·수집 시각 기록", "정적 프로토타입에서 확인 불가", "운영 검증 필요"],
        ["V-03", "스펙 완료율", "전체 자식 관계 기준 계산", "마법가마솥 7/14, 제련소 4/10 확인", "통과"],
        ["V-04", "담당별 현재 업무", "완료 제외 현재 뷰 기준", "15명·77개 연결, 토글 상세 확인", "통과"],
        ["V-05", "GPT 사실 기반 요약", "근거 없는 상태·기한·담당 생성 금지", "요약 계약은 정의됨. 실제 모델·스키마 검증 전", "부분 통과"],
        ["V-06", "출처 충돌", "자동 덮어쓰기 없이 출처별 주장 표시", "마법가마솥 7/31 vs 약 8/7 표시", "통과"],
        ["V-07", "프로젝트 스펙 상세", "토글로 자식 일감·상태·기간 확인", "양 프로젝트 우선 6개 스펙 확인", "통과"],
        ["V-08", "담당자 상세", "모든 사람의 실제 일감 확인", "15개 토글, Kady 15개 포함 확인", "통과"],
        ["V-09", "대표 브리핑 순서", "결정 → 막힘 → 변경 순", "화면 구조 반영", "통과"],
        ["V-10", "삭제·오래된 데이터", "조용히 정상 처리하지 않고 오류 표시", "삭제 상태 불일치 사례 표시", "통과"],
        ["V-11", "근거 추적", "요약에서 Notion·Slack 원문 이동", "사례 링크 제공. 전체 항목 자동 연결은 미검증", "조건부 통과"],
        ["V-12", "반응형·상호작용", "390·768·1440px에서 토글·레이아웃 정상", "세 폭에서 오버플로·JS 오류 없음", "통과"],
        ["V-13", "의존관계 커버리지", "누락률을 화면에 표시", "실제 등록률 지표 없음", "운영 검증 필요"],
        ["V-14", "회의록 포함", "결정·합의가 프로젝트 요약에 반영", "회의록 실수집 데이터 미확인", "운영 검증 필요"],
    ]
    add_table(doc, ["ID", "검증 항목", "기대 결과", "현재 확인", "상태"], test_rows, [0.55, 1.25, 1.75, 2.05, 0.9], 8.0)

    add_heading(doc, "5.1 화면 테스트 증거", 2)
    add_bullet(doc, "뷰포트: 390px, 768px, 1440px")
    add_bullet(doc, "포지 앤 포춘: 우선 스펙 6개, 전체 자식 상세 67개, 완료율 미터 6개")
    add_bullet(doc, "피자레디: 우선 스펙 6개, 전체 자식 상세 35개, 완료율 미터 6개")
    add_bullet(doc, "담당별 흐름: 15명 토글, Kady 토글에서 현재 일감 15개")
    add_bullet(doc, "문서 너비와 뷰포트 너비 일치, 실행 중 JavaScript 오류 없음")

    add_page_break(doc)
    add_heading(doc, "6. 운영 수용 기준", 1)
    add_heading(doc, "6.1 구현 완료 전 필수 조건", 2)
    acceptance = [
        "매일 08:00 KST 수집 성공 여부, 마지막 성공 시각, 출처별 최신 시각을 화면에 표시한다.",
        "수집 일부 실패 시 ‘어디까지 믿어도 되는지’를 프로젝트별로 명시한다.",
        "핵심 스펙 생성 시 기본 의존관계를 자동 제안하고, 의존관계 등록률을 상시 노출한다.",
        "상태 판정은 Notion 원장을 사용하고 Slack·회의록의 자연어 추출은 원문 발췌·링크·미확인 상태를 보존한다.",
        "핵심 필드(범위·기한·담당·상태) 충돌만 확인 항목으로 만들고 나머지는 로그로 남긴다.",
        "GPT 출력은 JSON Schema로 검증하고 실패 시 규칙 기반 사실 목록으로 대체한다.",
        "요약 문장마다 최소 1개 이상의 근거 출처를 연결한다.",
        "개인 탭의 공개 범위와 사용 목적을 문서화하고 성과 평가 용도로 사용하지 않는다.",
    ]
    for item in acceptance:
        add_bullet(doc, "☐ " + item)

    add_heading(doc, "6.2 1주 파일럿 측정 지표", 2)
    metric_rows = [
        ["의존관계 등록률", "핵심 스펙 중 의존관계가 있는 비율", "60% 미만이면 판정 신뢰 불가"],
        ["결정 큐 정밀도", "표시 안건 중 실제 결정이 필요했던 비율", "목표 80% 이상"],
        ["병목 정밀도", "‘대기’ 표시 중 실제 대기였던 비율", "오탐을 사례별 분류"],
        ["확인 부채", "생성 확인 건수 대비 처리 건수", "미처리 순증가 금지"],
        ["추가 탐색 횟수", "대시보드 후 Slack·Notion 재확인 횟수", "감소 추세"],
        ["데이터 신선도", "08:00 수집 성공·출처별 지연", "실패·지연을 숨기지 않음"],
    ]
    add_table(doc, ["지표", "측정 방법", "판정 기준"], metric_rows, [1.35, 3.0, 2.15], 9.2)

    add_page_break(doc)
    add_heading(doc, "7. 알려진 한계와 잔여 위험", 1)
    risks = [
        ["R-01", "높음", "현재 산출물은 정적 데이터 스냅샷이다. 자동 수집·재시도·실패 알림이 없다.", "운영 파이프라인 통합 테스트"],
        ["R-02", "높음", "Slack 자연어를 구조화하는 추출 단계의 오탐·미탐이 실제 데이터로 검증되지 않았다.", "원문 발췌·미확인 상태·샘플 수기 채점"],
        ["R-03", "높음", "의존관계 미등록 시 ‘병목 없음’으로 보일 수 있다.", "등록률 노출·기본 템플릿"],
        ["R-04", "중간", "프로토타입은 각 프로젝트의 우선 스펙 6개를 상세 노출한다. 전체 52개 탐색성은 별도 확인이 필요하다.", "검색·필터·전체 보기 수용성 테스트"],
        ["R-05", "중간", "회의록의 실제 범위와 최신성이 검증되지 않았다.", "회의록 원본 연결과 수집 커버리지 표시"],
        ["R-06", "중간", "담당별 수량은 현재 뷰 기준이라 완료 이력과 총 투입량을 의미하지 않는다.", "화면에 기준 문구 상시 표시"],
        ["R-07", "낮음", "대표의 결정을 저장하지 않는 읽기 전용 구조다.", "의도된 비기능. 원본 시스템에서 후속 반영"],
    ]
    add_table(doc, ["ID", "등급", "위험", "대응"], risks, [0.55, 0.65, 3.65, 1.65], 8.6)
    add_callout(
        doc,
        "가장 위험한 실패",
        "잘못된 경고 한두 건보다 더 위험한 것은 의존관계·수집 실패가 누락됐는데도 ‘막힌 것 없음’으로 보이는 조용한 거짓 안심이다. 커버리지와 데이터 신선도를 첫 화면에 함께 보여줘야 한다.",
        LIGHT_RED,
        RED,
    )

    add_heading(doc, "7.1 검증 후 권고", 2)
    add_bullet(doc, "프로토타입 구조는 유지하되, 실제 구현의 첫 완료 기준을 ‘자동 수집 성공 + 근거 링크 + 열화 표시’로 정한다.")
    add_bullet(doc, "GPT 문장 품질보다 먼저 입력 커버리지와 충돌 탐지 정확도를 파일럿에서 측정한다.")
    add_bullet(doc, "대표가 대시보드만 본 뒤 추가 탐색 없이 올바른 다음 행동을 정했는지 사례 단위로 채점한다.")
    add_bullet(doc, "1주 파일럿에서 결정 큐 정밀도 80% 이상과 확인 부채 비증가를 달성한 뒤 대상 프로젝트를 확대한다.")

    add_page_break(doc)
    add_heading(doc, "8. 검토자 확인 및 승인", 1)
    add_body(doc, "아래 항목은 실제 운영 데이터가 연결된 뒤 검토자가 직접 확인한다.")
    sign_rows = [
        ["검토자 / 역할", ""],
        ["검토 일시", ""],
        ["검증 환경", ""],
        ["판정", "☐ 통과   ☐ 조건부 통과   ☐ 재검증   ☐ 실패"],
        ["주요 발견사항", ""],
        ["필수 수정사항", ""],
        ["서명", ""],
    ]
    add_table(doc, ["항목", "기록"], sign_rows, [1.5, 5.0], 10)

    add_heading(doc, "8.1 검토자 최종 질문", 2)
    final_questions = [
        "첫 화면 30초 안에 오늘 확인·결정할 안건을 말할 수 있었는가?",
        "마법가마솥이 ‘50% 완료’여도 왜 일정 확인이 필요한지 설명할 수 있었는가?",
        "포지 앤 포춘 07/13 일감에서 완료율보다 확인 요청 큐가 중요하다는 점이 보였는가?",
        "담당별 흐름을 개인 평가가 아니라 팀 대기 원인 탐색에 사용했는가?",
        "모든 핵심 주장에 원문 출처 또는 ‘정보 부족’ 표시가 있었는가?",
        "수집 실패·의존관계 누락이 ‘정상’으로 보이지 않았는가?",
    ]
    for item in final_questions:
        add_bullet(doc, "☐ " + item)

    add_heading(doc, "8.2 참조 위치", 2)
    refs = [
        ("프로토타입 HTML", "/Users/molip/.gstack/projects/newtaskdashboard/designs/decision-dashboard-20260714/finalized.html", None),
        ("프로젝트 리스트", None, "https://app.notion.com/p/27eb4a4650038016a5fef8ce4bff328c?v=27eb4a465003809599ec000c27d45503"),
        ("피자레디 현재 뷰", None, "https://app.notion.com/p/265b4a465003807b80adf815349c4562?v=265b4a4650038074a578000c37538021"),
        ("포지 앤 포춘 현재 뷰", None, "https://app.notion.com/p/306b4a4650038140aabbdeb52d7f05ca?v=306b4a465003810b8994000c957221cc"),
        ("마법가마솥 Notion", None, "https://app.notion.com/p/35eb4a46500380ddb0dde7d8fc52d438"),
        ("마법가마솥 Slack", None, "https://molip0225.slack.com/archives/C05FT5754SK/p1781258173365049?thread_ts=1781258173.365049&cid=C05FT5754SK"),
    ]
    for label, local, url in refs:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(label + ": ")
        set_run_font(r, 10, True, TEXT)
        if url:
            add_hyperlink(p, "원문 열기", url)
        else:
            r2 = p.add_run(local)
            set_run_font(r2, 9.5, False, MUTED)

    for sec in doc.sections:
        sec.top_margin = Inches(0.75)
        sec.bottom_margin = Inches(0.75)
        sec.left_margin = Inches(1.0)
        sec.right_margin = Inches(1.0)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_document()
