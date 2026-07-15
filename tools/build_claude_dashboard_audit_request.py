from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from build_dashboard_validation_doc import (
    AMBER,
    BLUE,
    DARK_BLUE,
    GREEN,
    LIGHT_AMBER,
    LIGHT_BLUE,
    LIGHT_GREEN,
    LIGHT_RED,
    MUTED,
    RED,
    TEXT,
    add_body,
    add_bullet,
    add_callout,
    add_heading,
    add_page_break,
    add_page_number,
    add_table,
    add_title,
    set_run_font,
)


ROOT = Path("/Volumes/PortableSSD/newtaskdashboard")
OUT = ROOT / "docs" / "Claude_업무현황_대시보드_구현검증_요청서_2026-07-15.docx"


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "AppleGothic"
    for attr in ("ascii", "hAnsi", "eastAsia"):
        normal._element.rPr.rFonts.set(qn(f"w:{attr}"), "AppleGothic")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Title", 26, DARK_BLUE, 0, 8),
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "AppleGothic"
        for attr in ("ascii", "hAnsi", "eastAsia"):
            style._element.rPr.rFonts.set(qn(f"w:{attr}"), "AppleGothic")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Bullet 2", "List Number"):
        style = doc.styles[name]
        style.font.name = "AppleGothic"
        for attr in ("ascii", "hAnsi", "eastAsia"):
            style._element.rPr.rFonts.set(qn(f"w:{attr}"), "AppleGothic")
        style.font.size = Pt(11)

    header = section.header.paragraphs[0]
    header.text = "업무현황 대시보드  |  독립 구현 검증 요청"
    header.alignment = 2
    for run in header.runs:
        set_run_font(run, 8.5, False, MUTED)
    add_page_number(section.footer.paragraphs[0])


def add_label_body(doc, label, body):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r1 = p.add_run(label + " ")
    set_run_font(r1, 11, True, DARK_BLUE)
    r2 = p.add_run(body)
    set_run_font(r2, 11, False, TEXT)
    return p


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    props = doc.core_properties
    props.title = "Claude 업무현황 대시보드 구현 검증 요청서"
    props.subject = "기능 구현·목적 적합성·누락 기능 독립 감사"
    props.author = "Codex"

    add_title(
        doc,
        "Claude용 업무현황 대시보드\n구현 검증 요청서",
        "기능이 실제로 동작하는지, 목적에 맞는지, 무엇을 더 만들거나 덜어내야 하는지 독립 감사",
    )
    add_table(
        doc,
        ["구분", "내용"],
        [
            ["검증 대상", "업무현황 의사결정 대시보드 프로토타입 및 관련 코드"],
            ["핵심 목적", "의사결정 지연과 팀 대기시간을 빠르게 발견·감소"],
            ["검증 방식", "화면 직접 조작 + 코드·데이터 흐름 확인 + 목적 적합성 평가"],
            ["요구 결과", "구현 여부, 결함, 누락 기능, 불필요 기능, 우선 보완안"],
        ],
        [1.3, 5.2],
        10,
    )
    add_callout(
        doc,
        "Claude에게 요구하는 결론",
        "설계 문서를 요약하지 말고 실제 구현물을 직접 검사하라. 화면에 보인다는 이유로 구현됐다고 판단하지 말며, 정적 샘플·하드코딩·실연동을 구분하라. 마지막에는 ‘이 대시보드만 보고 대표가 올바른 다음 행동을 정할 수 있는가’를 답하라.",
        LIGHT_AMBER,
        AMBER,
    )
    add_heading(doc, "검증 자료", 2)
    add_label_body(doc, "프로토타입", "/Users/molip/.gstack/projects/newtaskdashboard/designs/decision-dashboard-20260714/finalized.html")
    add_label_body(doc, "기존 기준서", "/Volumes/PortableSSD/newtaskdashboard/docs/업무현황_대시보드_검증서_2026-07-15.docx")
    add_body(doc, "기존 기준서의 ‘통과’ 판정은 정답이 아니다. 요구사항과 사례 데이터의 참고 자료로만 사용하고, 실제 구현 증거를 다시 수집한다.")

    add_page_break(doc)
    add_heading(doc, "1. 검증 임무", 1)
    add_body(doc, "당신은 제품 책임자·QA 리드·데이터 신뢰성 검토자의 역할을 동시에 맡은 독립 감사자다. 아래 네 가지를 증거로 판정하라.")
    mission_rows = [
        ["A. 기능 구현", "화면 요소가 실제로 동작하며 코드·데이터 흐름까지 연결되어 있는가?"],
        ["B. 목적 적합성", "대표가 프로젝트 진행·지연·병목·결정 필요를 빠르게 이해할 수 있는가?"],
        ["C. 데이터 신뢰", "Notion·Slack·회의록·GPT가 사실과 출처를 보존하며 오류를 숨기지 않는가?"],
        ["D. 범위 적정성", "빠진 필수 기능, 과한 기능, 중복 기능, 혼란을 만드는 표현은 무엇인가?"],
    ]
    add_table(doc, ["검증 축", "반드시 답할 질문"], mission_rows, [1.45, 5.05], 9.7)

    add_heading(doc, "1.1 절대 전제", 2)
    premises = [
        "대시보드는 읽기 전용이다. 대표의 결정을 저장하지 않는 것은 결함이 아니다.",
        "프로젝트 건강상태는 완료율만으로 판단하지 않는다. 스펙 확정, 의존관계, 기한, 후속 작업 영향을 함께 본다.",
        "Notion·Slack·회의록이 충돌하면 시스템이 최신 출처를 추측하거나 덮어쓰지 않고 각 주장을 보여준다.",
        "Slack 운영 신호는 별도 메뉴가 아니라 프로젝트 요약 또는 ‘확인 필요’에 포함한다.",
        "첫 화면 우선순위는 ‘오늘 결정할 것 → 현재 막힌 것 → 어제와 달라진 것’이다.",
        "담당별 흐름은 개인 성과평가가 아니라 팀 대기·의존성·미배정 집중을 찾는 화면이다.",
        "프로젝트는 프로젝트 리스트에서 ‘요약’이 체크된 대상을 사용한다.",
        "최상단 일감은 스펙이며, 스펙의 자식 일감에서 실제 파트별 작업을 확인한다.",
    ]
    for item in premises:
        add_bullet(doc, item)

    add_heading(doc, "1.2 검증 시 금지", 2)
    for item in [
        "설계서 문구를 그대로 반복하며 ‘구현됨’으로 판정하지 말 것",
        "정적 HTML에 표시된 시각을 실제 매일 수집 결과로 오인하지 말 것",
        "완료율이 높다는 이유로 정상 프로젝트라 판정하지 말 것",
        "의존관계가 없다는 이유로 병목이 없다고 단정하지 말 것",
        "추가 기능을 많이 제안하는 것을 좋은 개선안으로 여기지 말 것",
    ]:
        add_bullet(doc, item)

    add_page_break(doc)
    add_heading(doc, "2. 실제 구현 점검 절차", 1)
    add_heading(doc, "2.1 구현 형태부터 분류", 2)
    for item in [
        "프로토타입, 실제 애플리케이션, 정적 스냅샷, 하드코딩 데이터의 경계를 먼저 밝힌다.",
        "각 화면 숫자와 문장이 어디에서 생성되는지 코드에서 추적한다.",
        "Notion·Slack·회의록 커넥터, 매일 08:00 스케줄러, 저장소, 규칙 엔진, GPT 호출이 실제로 존재하는지 확인한다.",
        "존재하지 않거나 실행 증거가 없으면 ‘미구현’ 또는 ‘검증 불가’로 표시한다.",
        "마지막 수집 시각, 부분 실패, 재시도, 오류 로그가 화면과 운영 구조에 있는지 확인한다.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "2.2 화면과 상호작용", 2)
    interaction_rows = [
        ["반응형", "390px, 768px, 1440px", "잘림·가로 스크롤·겹침·터치 영역"],
        ["탭", "대표 브리핑·프로젝트·담당별 흐름·확인 필요", "모든 탭 이동과 상태 유지"],
        ["프로젝트", "요약 체크된 모든 프로젝트", "스펙 토글, 완료율, 자식 일감, 근거 링크"],
        ["담당별", "현재 뷰의 모든 개인", "토글, 실제 일감, 프로젝트·상위 스펙·상태·기간"],
        ["충돌", "Notion·Slack·회의록", "출처별 주장, 원문 링크, 자동 덮어쓰기 금지"],
        ["오류", "누락·삭제·빈값·링크 실패", "침묵하지 않고 열화·확인 필요 표시"],
    ]
    add_table(doc, ["대상", "조작", "확인 사항"], interaction_rows, [1.0, 2.05, 3.45], 9.2)
    add_body(doc, "각 실패에는 재현 절차, 예상 결과, 실제 결과, 화면 캡처 또는 코드 위치를 남긴다. 클릭해 보지 않은 기능은 통과 처리하지 않는다.")

    add_heading(doc, "2.3 샘플 데이터 교차검증", 2)
    add_callout(
        doc,
        "반드시 확인할 사례",
        "피자레디 ‘마법가마솥’: 전체 자식 14개, 완료 7개(50%), Notion 7/31과 Slack 약 8/7의 일정 충돌, SP61 제작 시작 대기. 이 사례를 보고 완료율과 건강상태를 분리해 해석할 수 있는지 검증한다.",
        LIGHT_BLUE,
        DARK_BLUE,
    )
    for item in [
        "포지 앤 포춘 ‘제련소·분류기’가 전체 자식 관계 10개 중 4개 완료로 40%인지 확인한다.",
        "포지 앤 포춘 ‘07/13 일감’의 완료율보다 확인 요청 큐가 병목으로 드러나는지 확인한다.",
        "담당별 흐름이 15명과 실제 현재 일감을 모두 보여주는지 확인한다.",
        "완료율은 부모 스펙의 전체 자식 관계, 담당별 업무량은 완료 제외 현재 뷰를 사용한다는 차이를 코드로 확인한다.",
    ]:
        add_bullet(doc, item)

    add_page_break(doc)
    add_heading(doc, "3. 목적 적합성 점검", 1)
    add_body(doc, "화면이 정상 동작해도 대표의 판단을 빠르게 만들지 못하면 목적에 실패한 것이다. 실제 화면만 본 상태에서 아래 질문에 30초·2분 테스트를 수행하라.")
    purpose_rows = [
        ["30초", "오늘 대표가 확인하거나 결정해야 할 것은 무엇인가?", "안건·영향·근거가 한 화면에 보임"],
        ["30초", "지금 팀을 가장 오래 기다리게 하는 병목은 무엇인가?", "대기 원인과 후속 영향이 명확함"],
        ["2분", "각 프로젝트가 어느 단계이고 얼마나 진행됐는가?", "단계·완료율·계획 대비 상태를 구분"],
        ["2분", "어제와 무엇이 달라졌는가?", "상태·기한·담당·충돌 변화가 근거와 표시"],
        ["2분", "누가 누구를 기다리는가?", "개인 수량이 아니라 의존 흐름을 설명"],
        ["2분", "이 화면을 어디까지 믿어도 되는가?", "수집 시각·커버리지·부분 실패가 보임"],
    ]
    add_table(doc, ["시간", "대표 질문", "합격 기준"], purpose_rows, [0.7, 3.05, 2.75], 9.1)

    add_heading(doc, "3.1 정보 구조 점검", 2)
    for item in [
        "경고 수가 많아도 무엇부터 볼지 우선순위가 명확한가?",
        "프로젝트 카드가 완료율·상태·위험·확인 필요를 혼동시키지 않는가?",
        "스펙 토글을 열기 전에도 현재 단계와 핵심 문제를 이해할 수 있는가?",
        "세부 일감은 충분하지만 대표 브리핑을 압도하지 않는가?",
        "근거 링크를 열지 않아도 판단 가능하고, 필요할 때 원문까지 추적 가능한가?",
        "빈 데이터와 병목 없음이 시각적으로 구분되는가?",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "3.2 GPT 요약 점검", 2)
    ai_rows = [
        ["사실성", "입력에 없는 상태·담당·기한·완료율을 만들지 않음"],
        ["출처", "핵심 주장마다 Notion·Slack·회의록 원문 또는 근거가 있음"],
        ["충돌", "서로 다른 주장을 함께 보이며 임의로 최신을 선택하지 않음"],
        ["부족", "근거가 없으면 ‘정보 부족’ 또는 ‘확인 필요’로 표현"],
        ["우선순위", "결정 → 막힘 → 변경 순으로 간결하게 요약"],
        ["실패 처리", "스키마 실패 시 GPT 문장을 버리고 사실 목록으로 대체"],
    ]
    add_table(doc, ["검증 항목", "합격 기준"], ai_rows, [1.35, 5.15], 9.4)
    add_body(doc, "프롬프트 문구만 존재하면 부분 구현이다. 실제 입력 샘플, 모델 출력, JSON Schema 검증, 실패 대체 경로를 실행하거나 코드로 확인해야 통과다.")

    add_page_break(doc)
    add_heading(doc, "4. 누락·과잉 기능 감사", 1)
    add_body(doc, "기능 제안은 많을수록 좋은 것이 아니다. 아래 분류로 나누고, 각각 대표의 판단 시간 또는 팀 대기시간에 미치는 영향을 설명하라.")
    priority_rows = [
        ["P0 필수", "없으면 잘못된 판단 또는 조용한 거짓 안심 발생", "출시 전 반드시 구현"],
        ["P1 중요", "판단 속도·신뢰를 크게 개선하지만 우회 가능", "파일럿 직후"],
        ["P2 선택", "편의·탐색성 개선", "사용 데이터 확인 후"],
        ["삭제/통합", "중복·혼란·운영 부채를 만드는 기능", "제거 또는 다른 화면에 흡수"],
    ]
    add_table(doc, ["분류", "판정 기준", "처리"], priority_rows, [1.05, 3.85, 1.6], 9.3)

    add_heading(doc, "4.1 우선 검토할 누락 후보", 2)
    candidates = [
        "매일 오전 자동 수집과 실패 알림, 마지막 성공 시각",
        "출처별 데이터 신선도와 부분 실패 시 열화 표시",
        "핵심 스펙 의존관계 등록률과 누락 경고",
        "기본 의존관계 템플릿과 예외 수정 방식",
        "계획 대비 판정을 위한 범위 확정 예정일·제작 완료 예정일",
        "GPT 입력·출력 스키마, 근거 링크, 실패 대체 템플릿",
        "Slack·회의록 추출 주장에 원문 발췌와 미확인 상태",
        "핵심 필드 충돌의 생성 임계값·담당·만료 규칙",
        "전체 스펙 검색·필터·우선 6개 외 탐색",
        "변경 이력의 기준 시점과 ‘어제와 달라진 것’ 계산",
    ]
    for item in candidates:
        add_bullet(doc, item)
    add_callout(
        doc,
        "중요",
        "위 목록을 그대로 요구사항으로 승인하지 말라. 실제 구현과 목적을 검토한 뒤 필요한 것만 P0·P1·P2로 남기고, 효과가 불명확한 것은 보류하라.",
        LIGHT_AMBER,
        AMBER,
    )

    add_heading(doc, "4.2 역방향 실패 테스트", 2)
    failure_rows = [
        ["의존관계 0건", "병목 없음으로 보이지 않고 커버리지 부족으로 경고하는가?"],
        ["Notion 성공·Slack 실패", "부분 수집과 신뢰 범위를 표시하는가?"],
        ["회의록만 일정 변경", "프로젝트 요약 또는 확인 필요에 근거와 반영되는가?"],
        ["삭제 일감이 진행 중", "완료율에 조용히 포함하지 않고 데이터 오류로 올리는가?"],
        ["모든 일감 미배정", "개인 탭에서 사라지지 않고 팀 큐 위험으로 보이는가?"],
        ["GPT 잘못된 기한 생성", "스키마·출처 검증이 차단하는가?"],
        ["충돌 100건", "확인 필요가 부채로 쌓이지 않도록 핵심 필드만 선별하는가?"],
    ]
    add_table(doc, ["상황", "확인 질문"], failure_rows, [1.55, 4.95], 9.2)

    add_page_break(doc)
    add_heading(doc, "5. Claude 최종 응답 형식", 1)
    add_body(doc, "최종 답변은 아래 순서를 지키고, 모든 판정에 화면·코드·원문 중 하나 이상의 증거를 붙인다.")
    output_rows = [
        ["1", "한 줄 결론", "통과 / 조건부 통과 / 실패 중 하나와 핵심 이유"],
        ["2", "구현 경계", "실연동·정적 샘플·하드코딩·미구현을 명확히 분류"],
        ["3", "점수표", "기능 구현, 목적 적합성, 데이터 신뢰, 병목 탐지, 설명 가능성, 운영 복원력 각 5점"],
        ["4", "구현 결함", "ID·심각도·재현·증거·영향·수정안·합격 기준"],
        ["5", "목적 부적합", "동작하지만 판단을 느리게 하거나 오해시키는 기능"],
        ["6", "누락 기능", "P0·P1·P2, 사용자 효과, 구현 복잡도, 권장 시점"],
        ["7", "삭제·통합", "불필요·중복·감시 오해·운영 부채를 만드는 요소"],
        ["8", "최우선 5개", "효과 대비 우선순위와 구체적 완료 조건"],
        ["9", "권장 v1 범위", "출시 전 필수·파일럿 후·제외 범위를 한 장으로 정리"],
        ["10", "재검증 계획", "자동화·실데이터·GPT·권한을 포함한 다음 테스트"],
    ]
    add_table(doc, ["순서", "섹션", "요구 내용"], output_rows, [0.55, 1.4, 4.55], 8.9)

    add_heading(doc, "5.1 판정 규칙", 2)
    verdict_rows = [
        ["통과", "핵심 기능이 실데이터로 작동하고, 대표가 추가 탐색 없이 올바른 다음 행동을 정할 수 있음"],
        ["조건부 통과", "화면 목적은 충족하지만 운영·데이터 신뢰의 핵심 조건이 남아 있음"],
        ["실패", "정적 시연 수준이거나 병목·충돌·수집 실패를 정상처럼 보이게 할 가능성이 큼"],
    ]
    add_table(doc, ["판정", "조건"], verdict_rows, [1.15, 5.35], 9.4)

    add_heading(doc, "5.2 발견사항 작성 예시", 2)
    example_rows = [
        ["F-01", "치명적", "의존관계 미등록 시 병목 없음으로 표시", "스펙 A 의존성 필드 제거 후 첫 화면", "커버리지 지표와 기본 템플릿"],
        ["F-02", "높음", "08:00 수집 문구는 있으나 스케줄러 코드 없음", "정적 HTML 상단 시각·코드 검색", "미구현으로 분류하고 운영 파이프라인 추가"],
    ]
    add_table(doc, ["ID", "심각도", "발견", "증거·재현", "권고"], example_rows, [0.55, 0.65, 1.8, 1.65, 1.85], 8.5)
    add_callout(
        doc,
        "최종 질문",
        "이 대시보드가 없을 때보다 대표의 판단이 실제로 빨라지는가? 빨라진다면 어떤 판단이 얼마나 줄어드는가? 빨라지지 않는다면 화면·데이터·운영 중 어디가 원인인가?",
        LIGHT_GREEN,
        GREEN,
    )

    add_page_break(doc)
    add_heading(doc, "6. Claude에 바로 전달할 실행 명령", 1)
    add_callout(
        doc,
        "복사 시작",
        "첨부된 ‘Claude 업무현황 대시보드 구현 검증 요청서’를 최상위 지침으로 사용하라. finalized.html을 브라우저에서 직접 열고 모든 탭·토글·링크를 조작하며, 관련 코드를 읽어 정적 샘플·하드코딩·실연동·미구현을 구분하라. 기존 검증서의 통과 판정을 신뢰하지 말고 요구사항 참고 자료로만 사용하라.",
        LIGHT_BLUE,
        DARK_BLUE,
    )
    command_items = [
        "먼저 이 시스템이 실제 애플리케이션인지 정적 프로토타입인지 한 줄로 판정하라.",
        "화면 기능을 직접 실행하고 실패는 재현 절차와 증거를 남겨라.",
        "의사결정 지연과 팀 대기시간 감소라는 목적에 각 기능이 기여하는지 평가하라.",
        "Notion·Slack·회의록·GPT·오전 자동 수집의 실제 구현 여부와 신뢰 경계를 확인하라.",
        "누락 기능은 P0·P1·P2, 불필요 기능은 삭제·통합으로 분류하라.",
        "완료율을 건강상태로 오인하거나 결정 저장 기능을 요구하지 마라.",
        "마법가마솥, 포지 앤 포춘, 담당별 전체 인원 사례를 실제 화면에서 교차검증하라.",
        "마지막에는 한 줄 결론, 점수표, 발견사항, 최우선 수정 5개, 권장 v1 범위, 재검증 계획을 제시하라.",
    ]
    for item in command_items:
        add_bullet(doc, item)
    add_callout(
        doc,
        "완료 조건",
        "‘화면이 보기 좋다’가 아니라, 대표가 30초 안에 오늘의 결정·병목을 말할 수 있고 그 판단을 출처와 데이터 신선도로 검증할 수 있을 때 완료로 본다.",
        LIGHT_RED,
        RED,
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
