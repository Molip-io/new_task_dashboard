from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from build_dashboard_validation_doc import (
    AMBER,
    BLUE,
    DARK_BLUE,
    GREEN,
    LIGHT_AMBER,
    LIGHT_BLUE,
    LIGHT_GREEN,
    LIGHT_RED,
    MUTED,
    RED,
    TEXT,
    add_body,
    add_bullet,
    add_callout,
    add_heading,
    add_page_break,
    add_page_number,
    add_table,
    add_title,
    set_run_font,
)


ROOT = Path("/Volumes/PortableSSD/newtaskdashboard")
OUT = ROOT / "docs" / "업무현황_대시보드_기능적합성_검토요청서_2026-07-15.docx"


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "AppleGothic"
    for attr in ("ascii", "hAnsi", "eastAsia"):
        normal._element.rPr.rFonts.set(qn(f"w:{attr}"), "AppleGothic")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Title", 26, DARK_BLUE, 0, 8),
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "AppleGothic"
        for attr in ("ascii", "hAnsi", "eastAsia"):
            style._element.rPr.rFonts.set(qn(f"w:{attr}"), "AppleGothic")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Bullet 2", "List Number"):
        style = doc.styles[name]
        style.font.name = "AppleGothic"
        for attr in ("ascii", "hAnsi", "eastAsia"):
            style._element.rPr.rFonts.set(qn(f"w:{attr}"), "AppleGothic")
        style.font.size = Pt(11)

    header = section.header.paragraphs[0]
    header.text = "업무현황 대시보드  |  기능 적합성 검토"
    header.alignment = 2
    for run in header.runs:
        set_run_font(run, 8.5, False, MUTED)
    add_page_number(section.footer.paragraphs[0])


def add_label_body(doc, label, body):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r1 = p.add_run(label + " ")
    set_run_font(r1, 11, True, DARK_BLUE)
    r2 = p.add_run(body)
    set_run_font(r2, 11, False, TEXT)
    return p


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    props = doc.core_properties
    props.title = "업무현황 대시보드 기능 적합성 검토 요청서"
    props.subject = "목적 적합성, 누락 기능, 보완 기능, 단순화 검토"
    props.author = "Codex"

    add_title(
        doc,
        "업무현황 대시보드\n기능 적합성 검토 요청서",
        "목적에 맞게 기능이 구성됐는지, 무엇을 추가·보완·통합·제외해야 하는지 검토",
    )
    add_table(
        doc,
        ["구분", "내용"],
        [
            ["검토 대상", "업무현황 의사결정 대시보드의 기능 구성과 정보 구조"],
            ["핵심 목적", "의사결정 지연과 병목을 발견해 팀 전체의 대기시간 감소"],
            ["주요 사용자", "프로젝트 전반을 빠르게 파악하고 개입해야 하는 대표·리더"],
            ["검토 결과", "유지·보완·추가·통합·제외 기능과 권장 v1 범위"],
        ],
        [1.3, 5.2],
        10,
    )
    add_callout(
        doc,
        "이 문서의 검토 범위",
        "구현 방식이나 실행 상태를 시험하는 문서가 아니다. 현재 기획된 기능이 업무현황 대시보드의 목적에 충분히 기여하는지, 더 필요한 기능과 보완점이 무엇인지 제품 관점에서 검토한다.",
        LIGHT_AMBER,
        AMBER,
    )
    add_heading(doc, "검토자가 답해야 할 네 가지", 2)
    for item in [
        "현재 기능만으로 프로젝트가 어떻게 진행되고 얼마나 진행됐는지 알 수 있는가?",
        "대표가 오늘 확인·결정해야 할 것과 팀을 막는 병목을 빠르게 찾을 수 있는가?",
        "빠진 필수 기능과 목적에 비해 과하거나 중복된 기능은 무엇인가?",
        "출시용 v1에 남길 최종 기능과 이후로 미룰 기능은 무엇인가?",
    ]:
        add_bullet(doc, item)

    add_page_break(doc)
    add_heading(doc, "1. 대시보드 목적과 성공 조건", 1)
    add_heading(doc, "1.1 핵심 목적", 2)
    add_body(doc, "이 시스템의 목적은 업무를 많이 기록하거나 개인별 생산성을 평가하는 것이 아니다. 여러 프로젝트의 진행 흐름과 계획 대비 상태를 이해하고, 의사결정 지연·파트 간 의존성·정보 충돌로 발생하는 팀 대기를 빠르게 발견하는 것이다.")
    purpose_rows = [
        ["진행 파악", "프로젝트가 어떤 단계이며 얼마나 진행됐는가?"],
        ["계획 대비", "목표와 주요 예정일에 비해 정상·앞섬·지연 중 무엇인가?"],
        ["결정 우선순위", "오늘 대표가 확인하거나 결정해야 할 것은 무엇인가?"],
        ["병목", "무엇이 누구의 후속 작업을 기다리게 하는가?"],
        ["변화", "어제와 비교해 상태·기한·담당·위험이 어떻게 달라졌는가?"],
        ["신뢰", "Notion·Slack·회의록 중 무엇을 근거로 어디까지 믿어도 되는가?"],
    ]
    add_table(doc, ["목적", "대시보드가 답해야 할 질문"], purpose_rows, [1.35, 5.15], 9.5)

    add_heading(doc, "1.2 성공 조건", 2)
    for item in [
        "대표가 첫 화면에서 오늘의 확인·결정 안건과 가장 큰 병목을 우선순위대로 이해한다.",
        "프로젝트 탭에서 현재 단계, 완료율, 계획 대비 상태, 핵심 스펙의 위험과 후속 영향을 함께 본다.",
        "담당별 흐름에서 사람의 수량이 아니라 누가 무엇을 기다리고 어떤 일감이 팀 대기를 만드는지 확인한다.",
        "출처가 충돌하거나 정보가 부족하면 정상처럼 보이지 않고 확인 필요 또는 정보 부족으로 드러난다.",
        "요약만으로 다음 행동을 정할 수 있고 필요할 때 원문 근거까지 추적할 수 있다.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "1.3 비목표", 2)
    for item in [
        "개인 성과평가·생산성 순위·감시",
        "모든 세부 업무를 대시보드 안에서 편집·관리",
        "대표가 내린 결정을 대시보드에 저장",
        "출처 충돌을 시스템이 임의로 최신 정보로 결정",
        "완료율 하나로 프로젝트 건강상태를 판정",
    ]:
        add_bullet(doc, item)

    add_page_break(doc)
    add_heading(doc, "2. 현재 기획된 기능", 1)
    current_rows = [
        ["대표 브리핑", "결정 → 막힘 → 변경 순으로 실행 우선순위 제시", "오늘 무엇을 볼지"],
        ["프로젝트 목록", "요약 체크된 프로젝트의 현재 상태와 주요 위험", "어디에 개입할지"],
        ["스펙 상세", "최상위 스펙과 전체 자식 일감, 담당, 상태, 기간", "실제 진행 내용"],
        ["완료율", "완료 자식 ÷ 전체 자식 관계", "진행 참고값"],
        ["건강상태", "스펙 확정·의존관계·기한·후속 영향", "정상·위험 판단"],
        ["담당별 흐름", "전체 개인의 현재 일감과 기다리는 관계", "개인·팀 병목"],
        ["확인 필요", "핵심 필드 충돌·정보 부족·데이터 오류", "잘못된 판단 방지"],
        ["GPT 요약", "세 출처의 근거를 묶어 짧은 상황 설명", "읽는 시간 감소"],
        ["오전 수집", "선택 프로젝트의 Notion·Slack·회의록을 매일 갱신", "정보 신선도"],
    ]
    add_table(doc, ["기능", "현재 역할", "대표 질문"], current_rows, [1.3, 3.55, 1.65], 8.8)

    add_heading(doc, "2.1 반드시 유지할 판정 원칙", 2)
    principles = [
        "완료율과 프로젝트 건강상태는 분리한다.",
        "프로젝트 상태는 핵심 스펙 확정, 의존관계, 후속 영향으로 판단한다.",
        "Notion·Slack·회의록을 함께 요약하되 충돌은 자동 해결하지 않는다.",
        "Slack 운영 신호는 독립 메뉴가 아니라 프로젝트 요약 또는 확인 필요에 포함한다.",
        "스펙 상세·완료율은 전체 자식 관계를, 담당별 현재 업무는 완료 제외 현재 뷰를 사용한다.",
        "담당별 화면은 개인 평가가 아니라 팀 대기와 미배정 집중을 찾는 데 사용한다.",
        "GPT는 상태·담당·기한을 추측하지 않고 출처가 있는 사실만 요약한다.",
    ]
    for item in principles:
        add_bullet(doc, item)

    add_heading(doc, "2.2 현재 기능의 핵심 검토 질문", 2)
    for item in [
        "각 기능이 대표의 판단 시간을 실제로 줄이는가, 아니면 정보량만 늘리는가?",
        "같은 정보를 두 화면에서 반복해 보여주는 중복은 없는가?",
        "완료율·건강상태·확인 필요가 서로 다른 의미로 명확히 구분되는가?",
        "세부 일감이 충분하면서도 첫 화면의 우선순위를 압도하지 않는가?",
        "개인 탭이 팀 병목보다 개인 업무량 비교로 읽힐 위험은 없는가?",
    ]:
        add_bullet(doc, item)

    add_page_break(doc)
    add_heading(doc, "3. 기능별 목적 적합성 검토", 1)
    review_rows = [
        ["대표 브리핑", "한 화면에서 다음 행동이 정해지는가?", "우선순위·영향·근거가 함께 있는가?"],
        ["프로젝트", "진행 단계와 계획 대비 상태가 구분되는가?", "완료율만 강조되지 않는가?"],
        ["스펙 상세", "세부 진행과 다음 파트 대기가 보이는가?", "상위 스펙의 모든 담당자가 혼란을 만들지 않는가?"],
        ["담당별", "누가 누구를 기다리는지 설명하는가?", "수량이 개인 평가처럼 보이지 않는가?"],
        ["확인 필요", "잘못된 판단 가능성을 빠르게 드러내는가?", "항목이 너무 많이 쌓이지 않는가?"],
        ["GPT 요약", "정보를 줄이면서 핵심 판단 근거를 보존하는가?", "원문 없는 추론을 하지 않는가?"],
        ["데이터 운영", "신선도와 정보 부족을 사용자가 이해하는가?", "갱신 실패가 정상처럼 보이지 않는가?"],
    ]
    add_table(doc, ["영역", "목적 적합성 질문", "보완 관점"], review_rows, [1.1, 2.75, 2.65], 8.7)

    add_heading(doc, "3.1 대표 브리핑에서 반드시 보여줄 내용", 2)
    briefing = [
        "오늘 확인·결정할 안건: 질문, 필요한 결정, 기한, 영향받는 프로젝트·스펙",
        "현재 막힌 것: 막힌 원인, 기다리는 파트·사람, 후속 영향, 대기 기간",
        "어제와 달라진 것: 상태·기한·담당·스펙 확정·출처 충돌의 변화",
        "데이터 신뢰: 마지막 수집 시각, 일부 출처 실패, 의존관계 커버리지",
    ]
    for item in briefing:
        add_bullet(doc, item)

    add_heading(doc, "3.2 프로젝트에서 반드시 보여줄 내용", 2)
    project = [
        "프로젝트 목표와 목표일, 현재 단계, 계획 대비 상태",
        "핵심 스펙별 완료율, 건강상태, 확정 여부, 다음 단계",
        "스펙의 전체 자식 일감과 파트별 담당·상태·기간",
        "현재 막힌 원인과 후속 작업에 미치는 영향",
        "Notion·Slack·회의록 통합 요약과 출처별 근거·충돌",
    ]
    for item in project:
        add_bullet(doc, item)

    add_heading(doc, "3.3 담당별 흐름에서 반드시 보여줄 내용", 2)
    people = [
        "전체 개인의 현재 일감 수와 실제 일감 목록",
        "각 일감의 프로젝트·상위 스펙·상태·기간",
        "본인이 기다리는 결정·산출물과 본인을 기다리는 후속 작업",
        "미배정 팀 큐와 특정 역할에 몰린 대기",
        "개인 성과평가가 아니라 병목 탐색용이라는 기준 문구",
    ]
    for item in people:
        add_bullet(doc, item)

    add_page_break(doc)
    add_heading(doc, "4. 추가·보완 기능 후보", 1)
    add_body(doc, "아래 항목은 자동으로 채택할 요구사항이 아니다. 현재 목적을 달성하는 데 실제로 필요한지 검토하고 필수·중요·선택·제외로 분류한다.")
    candidate_rows = [
        ["계획 대비 기준선", "범위 확정 예정일·제작 완료 예정일", "막판 지연 발견 방지"],
        ["의존관계 커버리지", "핵심 스펙 중 관계 등록률과 누락 경고", "병목 없음의 거짓 안심 방지"],
        ["변경 기준", "어제 스냅샷과 상태·기한·담당 변화 비교", "변경 요약의 객관성"],
        ["데이터 신선도", "출처별 마지막 성공 시각과 부분 실패", "어디까지 믿을지 판단"],
        ["병목 영향도", "기다리는 인원·일감·기간·목표일 영향", "개입 우선순위 결정"],
        ["미배정 팀 큐", "개인에게 귀속되지 않은 일감의 집중", "숨은 팀 대기 발견"],
        ["충돌 처리 규칙", "핵심 필드만 생성, 담당·기한·만료", "확인 필요 부채 통제"],
        ["GPT 근거·대체", "문장별 근거와 실패 시 사실 목록", "요약 신뢰 확보"],
        ["전체 스펙 탐색", "우선 스펙 외 검색·필터·전체 보기", "누락 없이 필요할 때 탐색"],
        ["프로젝트 목표 요약", "목표·완료 조건·핵심 마일스톤", "진척 수치의 의미 부여"],
    ]
    add_table(doc, ["후보 기능", "내용", "기대 효과"], candidate_rows, [1.5, 3.15, 1.85], 8.5)

    add_heading(doc, "4.1 우선순위 판정 기준", 2)
    priority_rows = [
        ["P0 필수", "없으면 잘못된 판단·조용한 거짓 안심이 발생", "v1 포함"],
        ["P1 중요", "판단 속도와 신뢰를 크게 높임", "파일럿 직후"],
        ["P2 선택", "탐색·편의 개선", "사용 데이터 확인 후"],
        ["통합", "기존 기능에 흡수하는 편이 단순함", "별도 메뉴·카드 금지"],
        ["제외", "목적 기여가 낮거나 감시·운영 부채를 만듦", "만들지 않음"],
    ]
    add_table(doc, ["분류", "판정 기준", "권장 처리"], priority_rows, [1.0, 4.0, 1.5], 9.1)
    add_heading(doc, "4.2 기능 시나리오로 판단", 2)
    scenarios = [
        ["마법가마솥", "완료율 50%이나 Notion 7/31·Slack 약 8/7 충돌, 제작 시작 대기", "완료율과 건강상태·확인 필요가 분리되는가?"],
        ["포지 07/13", "완료율 3%, 확인 요청 17개", "낮은 완료율보다 확인 큐 병목이 우선으로 보이는가?"],
        ["담당 15개", "한 개인에게 현재 일감 15개", "업무량 평가가 아니라 팀 대기 원인으로 해석되는가?"],
        ["의존관계 없음", "관계가 입력되지 않은 핵심 스펙", "병목 없음이 아니라 정보 부족으로 보이는가?"],
    ]
    add_table(doc, ["상황", "관찰 사실", "필요한 기능 판단"], scenarios, [1.1, 3.2, 2.2], 8.0)

    add_heading(doc, "5. 검토 결과 작성 형식", 1)
    add_body(doc, "최종 검토는 기능을 많이 제안하는 것이 아니라 목적에 필요한 최소한의 완성된 구성을 만드는 데 초점을 둔다.")
    result_rows = [
        ["1", "한 줄 결론", "현재 기능 구성이 목적에 적합 / 부분 적합 / 재구성 필요"],
        ["2", "목적 충족도", "진행 파악·계획 대비·결정·병목·변화·신뢰 각 5점"],
        ["3", "현재 기능 판정", "기능별 유지·보완·통합·제외와 이유"],
        ["4", "누락 기능", "P0·P1·P2와 사용자 효과"],
        ["5", "혼란 요소", "오해·중복·정보 과다·감시 인식을 만드는 요소"],
        ["6", "첫 화면 권고", "보여줄 순서와 각 항목의 최소 정보"],
        ["7", "권장 v1", "출시 전 필수 기능의 최종 목록"],
        ["8", "이후 범위", "파일럿 후 검토할 기능과 제외할 기능"],
        ["9", "성공 지표", "판단 시간·추가 탐색·병목 정확도·확인 부채"],
    ]
    add_table(doc, ["순서", "결과 항목", "요구 내용"], result_rows, [0.55, 1.45, 4.5], 8.9)

    add_heading(doc, "5.1 현재 기능 판정 표", 2)
    decision_rows = [
        ["유지", "목적에 직접 기여하며 현재 위치가 적절함"],
        ["보완", "필요하지만 정보·기준·표현이 부족함"],
        ["통합", "필요하지만 별도 기능보다 기존 화면에 흡수해야 함"],
        ["제외", "목적 기여가 낮거나 부작용·운영 부채가 큼"],
        ["추가", "현재 기능으로는 핵심 판단 질문에 답할 수 없음"],
    ]
    add_table(doc, ["판정", "의미"], decision_rows, [1.0, 5.5], 9.5)

    add_heading(doc, "5.2 성공 지표 권고", 2)
    metric_rows = [
        ["첫 판단 시간", "첫 화면에서 오늘 결정·병목을 말하기까지 걸린 시간"],
        ["추가 탐색", "대시보드 확인 후 Notion·Slack을 다시 찾은 횟수"],
        ["결정 큐 정밀도", "표시 안건 중 실제 확인·결정이 필요했던 비율"],
        ["병목 정밀도", "대기 표시 중 실제 후속 작업이 기다리고 있던 비율"],
        ["확인 부채", "새로 생성된 확인 필요 대비 처리된 항목 수"],
        ["정보 커버리지", "핵심 스펙의 의존관계·기한·담당 등록률"],
    ]
    add_table(doc, ["지표", "정의"], metric_rows, [1.55, 4.95], 9.2)

    add_page_break(doc)
    add_heading(doc, "6. 검토자에게 전달할 요청문", 1)
    add_callout(
        doc,
        "검토 요청",
        "이 문서를 기준으로 현재 업무현황 대시보드의 기능 구성이 목적에 적합한지 검토해 주세요. 구현 상태나 화면 동작이 아니라, 대표가 프로젝트 진행·계획 대비·결정 필요·병목·변화·데이터 신뢰를 빠르게 이해하기에 기능이 충분한지를 평가해 주세요.",
        LIGHT_BLUE,
        DARK_BLUE,
    )
    request_items = [
        "현재 기능마다 유지·보완·통합·제외 중 하나를 판정하고 이유를 설명해 주세요.",
        "빠진 기능은 P0·P1·P2로 나누고 대표의 판단 또는 팀 대기시간에 미치는 효과를 적어 주세요.",
        "정보가 많아져 오히려 판단을 늦추거나 개인 감시로 오해될 요소를 찾아 주세요.",
        "완료율과 건강상태, 프로젝트 요약과 확인 필요, 담당별 업무량과 병목이 명확히 구분되는지 검토해 주세요.",
        "첫 화면을 ‘오늘 결정할 것 → 현재 막힌 것 → 어제와 달라진 것’으로 구성하는 것이 충분한지 보완안을 제안해 주세요.",
        "Notion·Slack·회의록·GPT를 합친 요약에서 반드시 보여줄 내용과 보여주지 말아야 할 내용을 정리해 주세요.",
        "마지막에는 출시용 v1의 최종 기능 목록, 파일럿 이후 기능, 제외 기능을 구분해 주세요.",
    ]
    for item in request_items:
        add_bullet(doc, item)
    add_callout(
        doc,
        "최종 판단 질문",
        "이 기능 구성만으로 대표가 프로젝트가 어떻게 진행되고 얼마나 진행됐는지 이해하고, 지금 개입해야 할 의사결정과 병목을 놓치지 않을 수 있는가? 부족하다면 가장 먼저 무엇을 보완해야 하는가?",
        LIGHT_GREEN,
        GREEN,
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
